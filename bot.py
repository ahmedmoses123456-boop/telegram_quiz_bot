import os
import re
import json
import uuid
import random
import psycopg2
from docx import Document
import openpyxl
from datetime import datetime

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters,
    CallbackQueryHandler,
    PollAnswerHandler,
)

TOKEN = os.getenv("BOT_TOKEN")
DATABASE_URL = os.getenv("DATABASE_URL")
BOT_USERNAME = os.getenv("BOT_USERNAME", "").strip()

temp_uploads = {}
quiz_creation_step = {}
user_sessions = {}

# ===================== DATABASE =====================

def get_db_connection():
    return psycopg2.connect(DATABASE_URL)


def init_db():
    conn = get_db_connection()
    cur = conn.cursor()

    # quizzes table
    cur.execute("""
    CREATE TABLE IF NOT EXISTS quizzes (
        quiz_id TEXT PRIMARY KEY,
        quiz_name TEXT NOT NULL,
        questions_json TEXT NOT NULL,
        time_per_question INTEGER DEFAULT 30
    )
    """)

    # quiz_results table (old schema safe)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS quiz_results (
        id SERIAL PRIMARY KEY,
        quiz_id TEXT NOT NULL,
        user_id BIGINT NOT NULL,
        score INTEGER NOT NULL,
        total_questions INTEGER NOT NULL,
        duration_seconds INTEGER NOT NULL,
        started_at TIMESTAMP NOT NULL,
        finished_at TIMESTAMP NOT NULL
    )
    """)

    # ADD COLUMN SAFELY (fix old DB)
    cur.execute("""
    ALTER TABLE quiz_results
    ADD COLUMN IF NOT EXISTS is_completed BOOLEAN DEFAULT TRUE
    """)

    conn.commit()
    cur.close()
    conn.close()


def save_quiz_to_db(quiz_id, quiz_name, questions, time_per_question):
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute(
        "INSERT INTO quizzes (quiz_id, quiz_name, questions_json, time_per_question) VALUES (%s, %s, %s, %s)",
        (quiz_id, quiz_name, json.dumps(questions, ensure_ascii=False), time_per_question)
    )

    conn.commit()
    cur.close()
    conn.close()


def load_quiz_from_db(quiz_id):
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute(
        "SELECT quiz_name, questions_json, time_per_question FROM quizzes WHERE quiz_id = %s",
        (quiz_id,)
    )
    row = cur.fetchone()

    cur.close()
    conn.close()

    if not row:
        return None

    quiz_name, questions_json, time_per_question = row
    questions = json.loads(questions_json)

    return {
        "quiz_name": quiz_name,
        "questions": questions,
        "time_per_question": time_per_question
    }


def save_result(quiz_id, user_id, score, total_questions, duration_seconds, started_at, finished_at, is_completed=True):
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("""
    INSERT INTO quiz_results
    (quiz_id, user_id, score, total_questions, duration_seconds, started_at, finished_at, is_completed)
    VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
    """, (quiz_id, user_id, score, total_questions, duration_seconds, started_at, finished_at, is_completed))

    conn.commit()
    cur.close()
    conn.close()


def get_rank_for_result(quiz_id, score, duration_seconds):
    conn = get_db_connection()
    cur = conn.cursor()

    cur.execute("SELECT COUNT(*) FROM quiz_results WHERE quiz_id = %s AND is_completed = TRUE", (quiz_id,))
    total_users = cur.fetchone()[0]

    cur.execute("""
    SELECT COUNT(*) FROM quiz_results
    WHERE quiz_id = %s
    AND is_completed = TRUE
    AND (
        score > %s
        OR (score = %s AND duration_seconds < %s)
    )
    """, (quiz_id, score, score, duration_seconds))

    better_count = cur.fetchone()[0]

    cur.close()
    conn.close()

    rank = better_count + 1
    return rank, total_users


# ===================== HELPERS =====================

def normalize_header(text: str):
    return str(text).strip().lower()


def safe_str(value):
    if value is None:
        return ""
    return str(value).strip()


def format_duration(seconds: int):
    minutes = seconds // 60
    sec = seconds % 60
    return f"{minutes}m {sec}s"


def shuffle_question_options(question):
    options = question["options"]
    correct_index = question["correct"]

    correct_option = options[correct_index]

    new_options = options.copy()
    random.shuffle(new_options)

    new_correct_index = new_options.index(correct_option)

    question["options"] = new_options
    question["correct"] = new_correct_index

    return question


def get_controls_keyboard(paused=False):
    if paused:
        return InlineKeyboardMarkup([
            [
                InlineKeyboardButton("â–¶ï¸ Resume Quiz", callback_data="RESUME"),
                InlineKeyboardButton("ğŸ›‘ Stop Quiz", callback_data="STOP")
            ]
        ])
    else:
        return InlineKeyboardMarkup([
            [
                InlineKeyboardButton("â¸ Pause Quiz", callback_data="PAUSE"),
                InlineKeyboardButton("ğŸ›‘ Stop Quiz", callback_data="STOP")
            ]
        ])


def get_finish_keyboard(has_wrong_questions: bool):
    buttons = []
    if has_wrong_questions:
        buttons.append([InlineKeyboardButton("ğŸ” Retry Wrong Questions", callback_data="RETRY_WRONG")])
    return InlineKeyboardMarkup(buttons) if buttons else None


# ===================== PARSERS =====================

def parse_docx_old_format(full_text: str):
    blocks = re.split(r"\n\s*---\s*\n", full_text)
    questions = []

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        q_match = re.search(r"Q:\s*(.+)", block)
        if not q_match:
            continue

        question_text = q_match.group(1).strip()

        tf_match = re.search(r"TYPE:\s*TF", block, re.IGNORECASE)

        explanation_match = re.search(r"EXPLANATION:\s*(.+)", block, re.DOTALL)
        explanation = explanation_match.group(1).strip() if explanation_match else "No explanation provided."

        answer_match = re.search(r"ANSWER:\s*(.+)", block)
        if not answer_match:
            continue
        answer_raw = answer_match.group(1).strip()

        if tf_match:
            options = ["True", "False"]
            correct = 0 if answer_raw.lower() == "true" else 1
        else:
            options = []
            for letter in ["A", "B", "C", "D"]:
                opt_match = re.search(rf"{letter}\)\s*(.+)", block)
                if opt_match:
                    options.append(opt_match.group(1).strip())

            if len(options) < 2:
                continue

            correct_letter = answer_raw.upper().strip()
            correct_map = {"A": 0, "B": 1, "C": 2, "D": 3}

            if correct_letter not in correct_map:
                continue

            correct = correct_map[correct_letter]

        questions.append({
            "question": question_text,
            "options": options,
            "correct": correct,
            "explanation": explanation
        })

    return questions


def parse_docx_table(doc: Document):
    questions = []

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        headers = [normalize_header(cell.text) for cell in rows[0].cells]

        if "type" not in headers or "question" not in headers or "correct" not in headers or "explanation" not in headers:
            continue

        def get_cell(row, col_name):
            if col_name not in headers:
                return ""
            idx = headers.index(col_name)
            if idx >= len(row.cells):
                return ""
            return safe_str(row.cells[idx].text)

        for r in rows[1:]:
            q_type = get_cell(r, "type").upper()
            q_text = get_cell(r, "question")
            correct_val = get_cell(r, "correct")
            explanation = get_cell(r, "explanation")

            if not q_text:
                continue

            if not explanation:
                explanation = "No explanation provided."

            if q_type == "TF":
                options = ["True", "False"]

                if correct_val.lower() == "true":
                    correct = 0
                elif correct_val.lower() == "false":
                    correct = 1
                else:
                    continue

                questions.append({
                    "question": q_text,
                    "options": options,
                    "correct": correct,
                    "explanation": explanation
                })

            else:
                a = get_cell(r, "a")
                b = get_cell(r, "b")
                c = get_cell(r, "c")
                d = get_cell(r, "d")

                options = [a, b, c, d]
                options = [opt.strip() for opt in options if opt.strip()]

                if len(options) < 2:
                    continue

                correct_letter = correct_val.upper().strip()
                correct_map = {"A": 0, "B": 1, "C": 2, "D": 3}

                if correct_letter not in correct_map:
                    continue

                correct = correct_map[correct_letter]

                if correct >= len(options):
                    continue

                questions.append({
                    "question": q_text,
                    "options": options,
                    "correct": correct,
                    "explanation": explanation
                })

    return questions


def parse_docx(file_path):
    doc = Document(file_path)

    table_questions = parse_docx_table(doc)
    if table_questions:
        return table_questions

    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    full_text = "\n".join(lines)
    old_questions = parse_docx_old_format(full_text)

    return old_questions


def parse_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active

    questions = []

    header_row = []
    for cell in sheet[1]:
        header_row.append(normalize_header(cell.value))

    required = ["type", "question", "correct", "explanation"]
    if not all(r in header_row for r in required):
        return []

    def get_col_index(col_name):
        return header_row.index(col_name) + 1

    type_col = get_col_index("type")
    question_col = get_col_index("question")
    correct_col = get_col_index("correct")
    explanation_col = get_col_index("explanation")

    a_col = header_row.index("a") + 1 if "a" in header_row else None
    b_col = header_row.index("b") + 1 if "b" in header_row else None
    c_col = header_row.index("c") + 1 if "c" in header_row else None
    d_col = header_row.index("d") + 1 if "d" in header_row else None

    for row in range(2, sheet.max_row + 1):
        q_type = safe_str(sheet.cell(row=row, column=type_col).value).upper()
        q_text = safe_str(sheet.cell(row=row, column=question_col).value)
        correct_val = safe_str(sheet.cell(row=row, column=correct_col).value)
        explanation = safe_str(sheet.cell(row=row, column=explanation_col).value)

        if not q_text:
            continue

        if not explanation:
            explanation = "No explanation provided."

        if q_type == "TF":
            options = ["True", "False"]

            if correct_val.lower() == "true":
                correct = 0
            elif correct_val.lower() == "false":
                correct = 1
            else:
                continue

            questions.append({
                "question": q_text,
                "options": options,
                "correct": correct,
                "explanation": explanation
            })

        else:
            a = safe_str(sheet.cell(row=row, column=a_col).value) if a_col else ""
            b = safe_str(sheet.cell(row=row, column=b_col).value) if b_col else ""
            c = safe_str(sheet.cell(row=row, column=c_col).value) if c_col else ""
            d = safe_str(sheet.cell(row=row, column=d_col).value) if d_col else ""

            options = [a, b, c, d]
            options = [opt.strip() for opt in options if opt.strip()]

            if len(options) < 2:
                continue

            correct_letter = correct_val.upper().strip()
            correct_map = {"A": 0, "B": 1, "C": 2, "D": 3}

            if correct_letter not in correct_map:
                continue

            correct = correct_map[correct_letter]

            if correct >= len(options):
                continue

            questions.append({
                "question": q_text,
                "options": options,
                "correct": correct,
                "explanation": explanation
            })

    return questions


# ===================== TIMEOUT =====================

async def question_timeout(context: ContextTypes.DEFAULT_TYPE):
    user_id = context.job.data["user_id"]
    question_index = context.job.data["question_index"]

    session = user_sessions.get(user_id)
    if not session:
        return

    if session.get("paused"):
        return

    if question_index in session["answered"]:
        return

    session["answered"].add(question_index)

    # Retry mode messages
    if session.get("mode") == "retry":
        q_text = session["questions"][question_index]["question"]
        session["retry_wrong_counter"][q_text] = session["retry_wrong_counter"].get(q_text, 0) + 1
        count = session["retry_wrong_counter"][q_text]

        if count == 1:
            await context.bot.send_message(session["chat_id"], "Ø®Ù„Ø§Øµ Ø±ÙƒØ² Ø¨Ù‚Ù‰ ÙŠØ§ Bro ğŸ˜­")
        elif count == 2:
            await context.bot.send_message(session["chat_id"], "Ø£Ù†Ø§ Ø¨Ø¯Ø£Øª Ø£Ù‚Ù„Ù‚ Ø¹Ù„ÙŠÙƒ ÙŠØ§ Broâ€¦ Ø±ÙƒØ² Ø¨Ø¬Ø¯ ğŸ˜­ğŸ’”")
        else:
            await context.bot.send_message(session["chat_id"], "ÙŠØ§ Ø¹Ù… Ø§Ù†Øª ØºÙ„Ø·Øª ÙÙŠ Ø§Ù„Ø£Ø³Ø¦Ù„Ø© Ø¯ÙŠ Ù£ Ù…Ø±Ø§Øª Ø±ÙƒØ² ğŸ˜­ğŸ”¥")

    await send_next_question(user_id, context)


# ===================== BOT FLOW =====================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    args = context.args

    if args:
        quiz_id = args[0].strip()

        quiz_data = load_quiz_from_db(quiz_id)
        if not quiz_data:
            await update.message.reply_text("âŒ Quiz not found.")
            return

        quiz_name = quiz_data["quiz_name"]

        keyboard = [[InlineKeyboardButton("â–¶ï¸ Start your quiz", callback_data=f"STARTQUIZ|{quiz_id}")]]
        await update.message.reply_text(
            f"ğŸ“Œ Quiz Name: {quiz_name}\n\nPress the button below to start ğŸ‘‡",
            reply_markup=InlineKeyboardMarkup(keyboard)
        )
        return

    await update.message.reply_text("Welcome!\nSend me a DOCX or XLSX file containing your quiz questions.")


async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    doc_file = update.message.document
    filename = doc_file.file_name.lower()

    if not (filename.endswith(".docx") or filename.endswith(".xlsx")):
        await update.message.reply_text("âŒ Please send a .docx or .xlsx file only.")
        return

    file = await doc_file.get_file()
    os.makedirs("downloads", exist_ok=True)

    path = f"downloads/{user_id}_{filename}"
    await file.download_to_drive(path)

    questions = parse_docx(path) if filename.endswith(".docx") else parse_xlsx(path)

    if not questions:
        await update.message.reply_text("âŒ No valid questions found.")
        return

    temp_uploads[user_id] = {"questions": questions}
    quiz_creation_step[user_id] = {"step": "waiting_name", "quiz_name": ""}

    await update.message.reply_text(
        f"âœ… File received!\nQuestions extracted: {len(questions)}\n\n"
        "ğŸ“ Now send me the Quiz Name."
    )


async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    text = update.message.text.strip()

    if user_id not in quiz_creation_step:
        return

    if quiz_creation_step[user_id]["step"] == "waiting_name":
        quiz_creation_step[user_id]["quiz_name"] = text
        quiz_creation_step[user_id]["step"] = "waiting_timer"
        await update.message.reply_text("â±ï¸ Enter time per question in seconds (5 - 600):")
        return

    if quiz_creation_step[user_id]["step"] == "waiting_timer":

        if not text.isdigit():
            await update.message.reply_text("âŒ Please enter a number.")
            return

        seconds = int(text)
        if seconds < 5 or seconds > 600:
            await update.message.reply_text("âŒ Please choose a time between 5 and 600 seconds.")
            return

        quiz_name = quiz_creation_step[user_id]["quiz_name"]
        questions = temp_uploads[user_id]["questions"]

        quiz_id = "q_" + uuid.uuid4().hex[:8]
        save_quiz_to_db(quiz_id, quiz_name, questions, seconds)

        del temp_uploads[user_id]
        del quiz_creation_step[user_id]

        link = f"https://t.me/{BOT_USERNAME}?start={quiz_id}"

        await update.message.reply_text(
            f"ğŸ‰ Quiz saved successfully!\n\n"
            f"ğŸ“Œ Quiz Name: {quiz_name}\n"
            f"â±ï¸ Time per question: {seconds} seconds\n"
            f"ğŸ§¾ Questions: {len(questions)}\n\n"
            f"ğŸ”— Share Link:\n{link}"
        )


async def start_quiz_button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    quiz_id = query.data.split("|")[1]

    quiz_data = load_quiz_from_db(quiz_id)
    if not quiz_data:
        await query.message.reply_text("âŒ Quiz not found.")
        return

    questions = quiz_data["questions"].copy()
    random.shuffle(questions)
    questions = [shuffle_question_options(q.copy()) for q in questions]

    user_id = query.from_user.id
    chat_id = query.message.chat_id

    user_sessions[user_id] = {
        "quiz_id": quiz_id,
        "chat_id": chat_id,
        "questions": questions,
        "index": 0,
        "score": 0,
        "time_per_question": quiz_data.get("time_per_question", 30),
        "started_at": datetime.utcnow(),
        "answered": set(),
        "poll_to_index": {},
        "wrong_questions": [],
        "paused": False,
        "mode": "main",
        "retry_wrong_counter": {},
        "retry_failed": False,
        "retry_mistakes": 0
    }

    await query.message.reply_text("âœ… Quiz started!")
    await send_next_question(user_id, context)


async def send_next_question(user_id: int, context: ContextTypes.DEFAULT_TYPE):
    session = user_sessions.get(user_id)
    if not session:
        return

    if session.get("paused"):
        return

    idx = session["index"]
    questions = session["questions"]
    chat_id = session["chat_id"]
    t = session["time_per_question"]

    if idx >= len(questions):
        await finish_quiz(user_id, context, stopped=False)
        return

    q = questions[idx]

    message = await context.bot.send_poll(
        chat_id=chat_id,
        question=f"Q{idx+1}: {q['question']}",
        options=q["options"],
        type="quiz",
        correct_option_id=q["correct"],
        explanation=f"ğŸ’¡ {q['explanation']}",
        is_anonymous=False,
        open_period=t
    )

    poll_id = message.poll.id
    session["poll_to_index"][poll_id] = idx

    await context.bot.send_message(
        chat_id=chat_id,
        text="âš™ï¸ Controls:",
        reply_markup=get_controls_keyboard(paused=False)
    )

    old_jobs = context.job_queue.get_jobs_by_name(f"timeout_{user_id}")
    for job in old_jobs:
        job.schedule_removal()

    context.job_queue.run_once(
        question_timeout,
        when=t + 1,
        data={"user_id": user_id, "question_index": idx},
        name=f"timeout_{user_id}"
    )

    session["index"] += 1


async def poll_answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    poll_id = update.poll_answer.poll_id
    user_id = update.poll_answer.user.id

    session = user_sessions.get(user_id)
    if not session:
        return

    if session.get("paused"):
    return

    if poll_id not in session["poll_to_index"]:
        return

    idx = session["poll_to_index"][poll_id]

    if idx in session["answered"]:
        return

    session["answered"].add(idx)

    selected = update.poll_answer.option_ids[0]
    correct = session["questions"][idx]["correct"]

    if selected == correct:
        session["score"] += 1
    else:
        session["wrong_questions"].append(session["questions"][idx])

        if session["mode"] == "retry":
            session["retry_failed"] = True
            session["retry_mistakes"] += 1

            if session["retry_mistakes"] == 1:
                await context.bot.send_message(session["chat_id"], "Ø®Ù„Ø§Øµ Ø±ÙƒØ² Ø¨Ù‚Ù‰ ÙŠØ§ Bro ğŸ˜­")
            elif session["retry_mistakes"] == 2:
                await context.bot.send_message(session["chat_id"], "Ø£Ù†Ø§ Ø¨Ø¯Ø£Øª Ø£Ù‚Ù„Ù‚ Ø¹Ù„ÙŠÙƒ ÙŠØ§ Broâ€¦ Ø±ÙƒØ² Ø¨Ø¬Ø¯ ğŸ˜­ğŸ’”")
            else:
                await context.bot.send_message(session["chat_id"], "ÙŠØ§ Ø¹Ù… Ø§Ù†Øª ØºÙ„Ø·Øª ÙÙŠ Ø§Ù„Ø£Ø³Ø¦Ù„Ø© Ø¯ÙŠ Ù£ Ù…Ø±Ø§Øª Ø±ÙƒØ² ğŸ˜­ğŸ”¥")

    # remove old timer jobs
    old_jobs = context.job_queue.get_jobs_by_name(f"timeout_{user_id}")
    for job in old_jobs:
        job.schedule_removal()

    await send_next_question(user_id, context)


async def finish_quiz(user_id: int, context: ContextTypes.DEFAULT_TYPE, stopped=False):
    session = user_sessions.get(user_id)
    if not session:
        return

    score = session["score"]
    total = len(session["questions"])

    if total == 0:
        percent = 0
    else:
        percent = round((score / total) * 100, 1)

    finished_at = datetime.utcnow()
    started_at = session["started_at"]
    duration_seconds = int((finished_at - started_at).total_seconds())

    quiz_id = session["quiz_id"]

    # MAIN QUIZ RESULT (RANK INCLUDED)
    if session["mode"] == "main":
        save_result(
            quiz_id=quiz_id,
            user_id=user_id,
            score=score,
            total_questions=total,
            duration_seconds=duration_seconds,
            started_at=started_at,
            finished_at=finished_at,
            is_completed=True
        )

        rank, total_users = get_rank_for_result(quiz_id, score, duration_seconds)

        msg = (
            f"ğŸ‰ Quiz Finished!\n\n"
            f"ğŸ† Score: {score}/{total}\n"
            f"ğŸ“Š Percentage: {percent}%\n"
            f"â±ï¸ Duration: {format_duration(duration_seconds)}\n\n"
            f"ğŸ¥‡ Your Rank: {rank} / {total_users}"
        )

        await context.bot.send_message(
            chat_id=session["chat_id"],
            text=msg,
            reply_markup=get_finish_keyboard(has_wrong_questions=len(session["wrong_questions"]) > 0)
        )

    # RETRY RESULT (NO RANK)
    else:
        if session.get("retry_failed") is False:
            await context.bot.send_message(session["chat_id"], "Ø¹Ø§Ø§Ø§Ø§Ø´ ÙŠØ§ ÙˆØ­Ø´ ğŸ”¥ğŸ‰ Ø§Ù†Øª ÙÙ„Ù„Øª Ø§Ù„Ø±ÙŠØªØ±ÙŠ Ø¨Ø¯ÙˆÙ† ØºÙ„Ø·!")

        await context.bot.send_message(
            chat_id=session["chat_id"],
            text=(
                f"ğŸ” Retry Finished!\n\n"
                f"ğŸ† Score: {score}/{total}\n"
                f"ğŸ“Š Percentage: {percent}%\n"
                f"â±ï¸ Duration: {format_duration(duration_seconds)}\n\n"
                f"âœ… Retry Ù„Ø§ ÙŠØ­Ø³Ø¨ ÙÙŠ Ø§Ù„Ø±Ø§Ù†Ùƒ."
            )
        )

        del user_sessions[user_id]
        return

    del user_sessions[user_id]


async def retry_wrong_questions(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    session = user_sessions.get(user_id)

    if not session:
        await query.message.reply_text("âŒ No active session found.")
        return

    wrong = session.get("wrong_questions", [])

    if not wrong:
        await query.message.reply_text("ğŸ‰ No wrong questions! Ø£Ù†Øª Ø¬Ø§Ù…Ø¯ ğŸ˜ğŸ”¥")
        return

    random.shuffle(wrong)
    wrong = [shuffle_question_options(q.copy()) for q in wrong]

    session["questions"] = wrong
    session["index"] = 0
    session["score"] = 0
    session["answered"] = set()
    session["poll_to_index"] = {}
    session["started_at"] = datetime.utcnow()
    session["paused"] = False
    session["mode"] = "retry"
    session["retry_wrong_counter"] = {}
    session["retry_failed"] = False
    session["retry_mistakes"] = 0

    await query.message.reply_text(f"ğŸ” Retry Started! Questions: {len(wrong)}")
    await send_next_question(user_id, context)


async def pause_quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    session = user_sessions.get(user_id)

    if not session:
        await query.message.reply_text("âŒ No active quiz.")
        return

    session["paused"] = True

    old_jobs = context.job_queue.get_jobs_by_name(f"timeout_{user_id}")
    for job in old_jobs:
        job.schedule_removal()

    await query.message.reply_text("â¸ Quiz Paused!", reply_markup=get_controls_keyboard(paused=True))


async def resume_quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    session = user_sessions.get(user_id)

    if not session:
        await query.message.reply_text("âŒ No active quiz.")
        return

    session["paused"] = False
    await query.message.reply_text("â–¶ï¸ Quiz Resumed!")
    await send_next_question(user_id, context)


async def stop_quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user_id = query.from_user.id
    session = user_sessions.get(user_id)

    if not session:
        await query.message.reply_text("âŒ No active quiz.")
        return

    old_jobs = context.job_queue.get_jobs_by_name(f"timeout_{user_id}")
    for job in old_jobs:
        job.schedule_removal()

    await finish_quiz(user_id, context, stopped=True)


async def callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    data = query.data

    if data.startswith("STARTQUIZ|"):
        await start_quiz_button(update, context)

    elif data == "RETRY_WRONG":
        await retry_wrong_questions(update, context)

    elif data == "PAUSE":
        await pause_quiz(update, context)

    elif data == "RESUME":
        await resume_quiz(update, context)

    elif data == "STOP":
        await stop_quiz(update, context)


def main():
    if not TOKEN:
        print("ERROR: BOT_TOKEN missing!")
        return

    if not DATABASE_URL:
        print("ERROR: DATABASE_URL missing!")
        return

    if not BOT_USERNAME:
        print("ERROR: BOT_USERNAME missing!")
        return

    init_db()

    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(CallbackQueryHandler(callback_handler))
    app.add_handler(PollAnswerHandler(poll_answer))

    print("Bot is running...")
    app.run_polling()


if __name__ == "__main__":
    main()
